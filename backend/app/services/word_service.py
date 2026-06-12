"""
Génération de documents Word (.docx) pour les devis/factures.
Utilise python-docx (pur Python).

Deux modèles visuels :
  - "moderne" (défaut) : en-tête vert #14532D, Calibri, lots #E3EDE6
  - "pro"             : fond blanc, anthracite #1F2937, Georgia, lots bleu acier #3B5573
"""
import base64
from datetime import date as date_class
from typing import Optional
from io import BytesIO
from app.models.quote import Devis


PAGE_W_CM = 17.0


def _fmt_date(iso_date: Optional[str]) -> str:
    if iso_date:
        try:
            y, m, d = iso_date.split("-")
            return f"{d}/{m}/{y}"
        except Exception:
            pass
    return date_class.today().strftime("%d/%m/%Y")


def _logo_dimensions_cm(logo_data: bytes) -> tuple[float, float]:
    """Calcule la largeur et hauteur du logo (cm) en respectant l'aspect ratio."""
    W_MAX, H_MAX = 4.0, 2.5
    try:
        from PIL import Image as PILImage
        img = PILImage.open(BytesIO(logo_data))
        iw, ih = img.size
        aspect = iw / ih
        if aspect >= W_MAX / H_MAX:
            return W_MAX, W_MAX / aspect
        else:
            return H_MAX * aspect, H_MAX
    except Exception:
        return W_MAX, 2.0


def generate_quote_docx(
    devis: Devis,
    document_type: str = "devis",
    with_tva: bool = True,
    document_date: Optional[str] = None,
) -> bytes:
    """Génère le .docx et retourne les bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx n'est pas installé. Lancez : pip install python-docx")

    modele   = (devis.modele or "moderne").lower()
    is_pro   = (modele == "pro")
    BODY_FONT = "Georgia" if is_pro else "Calibri"

    doc_label = "FACTURE" if document_type == "facture" else "DEVIS"
    doc_date  = _fmt_date(document_date)

    def _fmt_money(amount: float) -> str:
        s = f"{amount:.2f}"
        int_part, dec_part = s.split(".")
        n, groups = len(int_part), []
        while n > 3:
            groups.insert(0, int_part[n - 3:n])
            n -= 3
        groups.insert(0, int_part[:n])
        return " ".join(groups) + "," + dec_part + " €"

    # ── Couleurs selon modèle ────────────────────────────────────────
    if is_pro:
        COL_ACCENT  = RGBColor(59, 85, 115)    # #3B5573 bleu acier
        COL_HEADER  = RGBColor(31, 41, 55)     # #1F2937 anthracite
        COL_INFO    = RGBColor(90, 99, 93)     # #5A635D gris
        COL_LOT_TXT = RGBColor(59, 85, 115)
        COL_LOT_BG  = None                     # pas de fond coloré
        HDR_BG      = None                     # pas de bandeau
        NET_BG      = None
        NET_TXT     = RGBColor(31, 41, 55)
        TTC_BG      = None
        TTC_TXT     = RGBColor(31, 41, 55)
    else:
        COL_ACCENT  = RGBColor(20, 83, 45)     # #14532D
        COL_HEADER  = RGBColor(255, 255, 255)  # blanc sur fond vert
        COL_INFO    = RGBColor(200, 230, 210)  # blanc cassé sur vert
        COL_LOT_TXT = RGBColor(20, 83, 45)
        COL_LOT_BG  = "E3EDE6"
        HDR_BG      = "14532D"
        NET_BG      = "047857"                 # vert foncé
        NET_TXT     = RGBColor(255, 255, 255)
        TTC_BG      = "14532D"
        TTC_TXT     = RGBColor(255, 255, 255)

    WHITE = RGBColor(255, 255, 255)
    BLACK = RGBColor(30, 30, 30)
    AMBER = RGBColor(180, 120, 0)

    # ── Helpers ──────────────────────────────────────────────────────

    def _set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _remove_cell_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    def _set_cell_border_bottom(cell, color: str = "1F2937", sz: str = "12"):
        """Filet uniquement en bas de la cellule (modèle pro)."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tcBorders.append(b)
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), sz)
        bot.set(qn("w:color"), color)
        tcBorders.append(bot)
        tcPr.append(tcBorders)

    TWIPS_PER_CM = 566.93

    def _fix_table(table, width_cm: float = PAGE_W_CM):
        w_twips = str(round(width_cm * TWIPS_PER_CM))
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        for tag in ("w:tblInd", "w:tblW", "w:tblLayout"):
            for old in tblPr.findall(qn(tag)):
                tblPr.remove(old)
        tblInd = OxmlElement("w:tblInd")
        tblInd.set(qn("w:w"), "0")
        tblInd.set(qn("w:type"), "dxa")
        tblPr.append(tblInd)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), w_twips)
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)

    def _zero_para_spacing(para):
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        pPr.append(spacing)

    def _run(para, text: str, size: int, color: RGBColor = None,
             bold: bool = False, italic: bool = False, font_name: str = None) -> None:
        r = para.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = font_name or BODY_FONT
        if color:
            r.font.color.rgb = color

    # ── Document ─────────────────────────────────────────────────────
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after  = Pt(0)

    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # ── EN-TÊTE : artisan (gauche) | label doc (droite) ─────────────
    ht = doc.add_table(rows=1, cols=2)
    ht.autofit = False
    ht.columns[0].width = Cm(10)
    ht.columns[1].width = Cm(7)
    ht.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fix_table(ht)

    for cell in ht.rows[0].cells:
        _remove_cell_borders(cell)
        if HDR_BG:
            _set_cell_bg(cell, HDR_BG)
        elif is_pro:
            # Pro : filet épais anthracite en bas de l'en-tête
            _set_cell_border_bottom(cell, color="1F2937", sz="16")

    lc = ht.cell(0, 0)

    # Logo artisan (optionnel)
    if devis.artisan.logo_base64:
        try:
            logo_data = base64.b64decode(devis.artisan.logo_base64)
            logo_w_cm, logo_h_cm = _logo_dimensions_cm(logo_data)
            p_logo = lc.paragraphs[0]
            _zero_para_spacing(p_logo)
            p_logo.add_run().add_picture(BytesIO(logo_data), width=Cm(logo_w_cm), height=Cm(logo_h_cm))
            p_name = lc.add_paragraph()
        except Exception:
            p_name = lc.paragraphs[0]
    else:
        p_name = lc.paragraphs[0]

    _zero_para_spacing(p_name)
    _run(p_name, devis.artisan.nom or "Votre Entreprise", size=14, color=COL_HEADER, bold=True)

    artisan_lines = [
        f"SIRET : {devis.artisan.siret}"   if devis.artisan.siret   else None,
        f"Adresse : {devis.artisan.adresse}" if devis.artisan.adresse else None,
        " ".join(filter(None, [devis.artisan.code_postal, devis.artisan.ville])) or None,
        f"Tel : {devis.artisan.telephone}" if devis.artisan.telephone else None,
        f"Mail : {devis.artisan.email}"    if devis.artisan.email    else None,
        f"Web : {devis.artisan.site_web}"  if devis.artisan.site_web else None,
    ]
    for line in artisan_lines:
        if line:
            p = lc.add_paragraph()
            _zero_para_spacing(p)
            _run(p, line, size=9, color=COL_INFO)

    rc = ht.cell(0, 1)
    p = rc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _zero_para_spacing(p)
    _run(p, doc_label, size=22, color=COL_HEADER if not is_pro else RGBColor(31, 41, 55), bold=True)

    p2 = rc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _zero_para_spacing(p2)
    _run(p2, f"Date : {doc_date}", size=9, color=COL_INFO if not is_pro else RGBColor(90, 99, 93))

    if document_type == "devis":
        p3 = rc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _zero_para_spacing(p3)
        _run(p3, "Valable 30 jours", size=9, color=COL_INFO if not is_pro else RGBColor(90, 99, 93))

    if devis.numero_document:
        num_label = "N Facture" if document_type == "facture" else "N Devis"
        p_num = rc.add_paragraph()
        p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _zero_para_spacing(p_num)
        _run(p_num, f"{num_label} : {devis.numero_document}", size=9,
             color=COL_HEADER if not is_pro else COL_ACCENT, bold=True)

    sep = doc.add_paragraph()
    _zero_para_spacing(sep)

    # ── CLIENT + CHANTIER ────────────────────────────────────────────
    it = doc.add_table(rows=1, cols=2)
    it.autofit = False
    it.columns[0].width = Cm(8.5)
    it.columns[1].width = Cm(8.5)
    it.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fix_table(it)

    cc = it.cell(0, 0)
    _set_cell_bg(cc, "EFF6FF")
    p = cc.paragraphs[0]
    _zero_para_spacing(p)
    _run(p, "CLIENT", size=8, color=COL_ACCENT, bold=True)
    p2 = cc.add_paragraph()
    _zero_para_spacing(p2)
    _run(p2, devis.client.nom or "—", size=10, bold=True)
    if devis.client.adresse:
        p3 = cc.add_paragraph()
        _zero_para_spacing(p3)
        _run(p3, devis.client.adresse, size=9, color=RGBColor(90, 99, 93))
    cp_ville_client = " ".join(filter(None, [devis.client.code_postal, devis.client.ville]))
    if cp_ville_client:
        p4 = cc.add_paragraph()
        _zero_para_spacing(p4)
        _run(p4, cp_ville_client, size=9, color=RGBColor(90, 99, 93))

    ch = it.cell(0, 1)
    _set_cell_bg(ch, "FFFBEB")
    p = ch.paragraphs[0]
    _zero_para_spacing(p)
    _run(p, "CHANTIER", size=8, color=AMBER, bold=True)
    p2 = ch.add_paragraph()
    _zero_para_spacing(p2)
    _run(p2, devis.chantier.description, size=9, color=BLACK)

    sep2 = doc.add_paragraph()
    _zero_para_spacing(sep2)

    # ── TABLEAU DES PRESTATIONS ──────────────────────────────────────
    # Colonnes : Prestation | Description | Qté | Unité | PU HT | [TVA] | Total HT
    if with_tva:
        headers    = ["Prestation", "Description", "Qte", "Unite", "PU HT", "TVA", "Total HT"]
        col_widths = [Cm(2.8), Cm(5.5), Cm(1.0), Cm(1.2), Cm(2.1), Cm(1.5), Cm(2.9)]  # T=17.0
        col_aligns = ([WD_ALIGN_PARAGRAPH.LEFT] * 2 +
                      [WD_ALIGN_PARAGRAPH.CENTER] * 2 +
                      [WD_ALIGN_PARAGRAPH.RIGHT] * 3)
    else:
        headers    = ["Prestation", "Description", "Qte", "Unite", "PU HT", "Total HT"]
        col_widths = [Cm(3.2), Cm(6.7), Cm(1.0), Cm(1.3), Cm(2.4), Cm(2.4)]  # T=17.0
        col_aligns = ([WD_ALIGN_PARAGRAPH.LEFT] * 2 +
                      [WD_ALIGN_PARAGRAPH.CENTER] * 2 +
                      [WD_ALIGN_PARAGRAPH.RIGHT] * 2)

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fix_table(tbl)
    for i, w in enumerate(col_widths):
        tbl.columns[i].width = w

    HDR_COLOR = "3B5573" if is_pro else "14532D"
    for i, (cell, h) in enumerate(zip(tbl.rows[0].cells, headers)):
        _set_cell_bg(cell, HDR_COLOR)
        p = cell.paragraphs[0]
        p.alignment = col_aligns[i]
        _zero_para_spacing(p)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = BODY_FONT
        r.font.color.rgb = WHITE

    # Lignes de prestation
    has_lots  = any(l.lot for l in devis.lignes)
    lot_groups: dict = {}
    for ligne in devis.lignes:
        key = ligne.lot or ""
        if key not in lot_groups:
            lot_groups[key] = []
        lot_groups[key].append(ligne)

    row_i = 0
    for lot_name, lot_lignes in lot_groups.items():

        # En-tête de LOT
        if has_lots and lot_name:
            lot_row = tbl.add_row()
            for cell in lot_row.cells:
                _remove_cell_borders(cell)
                if COL_LOT_BG:
                    _set_cell_bg(cell, COL_LOT_BG)
                elif is_pro:
                    _set_cell_border_bottom(cell, color="3B5573", sz="6")
            p = lot_row.cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _zero_para_spacing(p)
            r = p.add_run(f"  {lot_name}")
            r.bold = True
            r.font.size = Pt(9)
            r.font.name = BODY_FONT
            r.font.color.rgb = COL_LOT_TXT
            for cell in lot_row.cells[1:]:
                _zero_para_spacing(cell.paragraphs[0])

        for ligne in lot_lignes:
            montant_ht = round(ligne.quantite * ligne.prix_unitaire_ht, 2)
            bg = "F8F9FA" if row_i % 2 == 1 else "FFFFFF"
            row = tbl.add_row()

            qty_str = (f"{ligne.quantite:g}")
            vals = (
                [ligne.poste, ligne.description,
                 qty_str,
                 ligne.unite,
                 _fmt_money(ligne.prix_unitaire_ht),
                 f"{ligne.tva_taux:.0f}%",
                 _fmt_money(montant_ht)]
                if with_tva else
                [ligne.poste, ligne.description,
                 qty_str,
                 ligne.unite,
                 _fmt_money(ligne.prix_unitaire_ht),
                 _fmt_money(montant_ht)]
            )

            for cell, val, align in zip(row.cells, vals, col_aligns):
                _set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                p.alignment = align
                _zero_para_spacing(p)
                r = p.add_run(val)
                r.font.size = Pt(9)
                r.font.name = BODY_FONT
            row_i += 1

        # Sous-total du LOT
        if has_lots and len(lot_groups) > 1:
            lot_ht = round(sum(l.quantite * l.prix_unitaire_ht for l in lot_lignes), 2)
            sub_row = tbl.add_row()
            n = len(sub_row.cells)
            sub_label = f"Sous-total {lot_name}" if lot_name else "Sous-total"
            for idx, cell in enumerate(sub_row.cells):
                _set_cell_bg(cell, "F1F5F9")
                p = cell.paragraphs[0]
                _zero_para_spacing(p)
                if idx == n - 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    r = p.add_run(sub_label)
                    r.italic = True
                    r.font.size = Pt(9)
                    r.font.name = BODY_FONT
                    r.font.color.rgb = RGBColor(90, 99, 93)
                elif idx == n - 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    r = p.add_run(_fmt_money(lot_ht))
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.name = BODY_FONT
                    r.font.color.rgb = BLACK

    sep3 = doc.add_paragraph()
    _zero_para_spacing(sep3)

    # ── TOTAUX ───────────────────────────────────────────────────────
    tot = doc.add_table(rows=0, cols=2)
    tot.autofit = False
    tot.columns[0].width = Cm(4)
    tot.columns[1].width = Cm(3)
    tot.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _fix_table(tot, width_cm=7.0)

    totaux      = devis.totaux
    has_remise  = (totaux.remise_ht or 0) > 0
    has_acompte = (devis.acompte or 0) > 0

    def _add_total_row(label: str, value: str, bold: bool = False,
                       bg_hex: str = "F8F9FA", txt_color: RGBColor = None):
        r = tot.add_row()
        for cell in r.cells:
            _set_cell_bg(cell, bg_hex)
        if is_pro and bg_hex in ("14532D", "047857"):
            # Pro n'utilise pas ces fonds — remplacer par filet épais
            bg_hex = "F8F9FA"
        txt = txt_color or BLACK
        for cell, text, align in [(r.cells[0], label, WD_ALIGN_PARAGRAPH.LEFT),
                                   (r.cells[1], value, WD_ALIGN_PARAGRAPH.RIGHT)]:
            p = cell.paragraphs[0]
            p.alignment = align
            _zero_para_spacing(p)
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(10)
            run.font.name = BODY_FONT
            if txt:
                run.font.color.rgb = txt

    def _add_total_accent(label: str, value: str):
        """Ligne TOTAL TTC ou NET À PAYER selon modèle."""
        r = tot.add_row()
        if is_pro:
            bg_hex = "F8F9FA"
            for cell in r.cells:
                _set_cell_bg(cell, bg_hex)
                _set_cell_border_bottom(cell, color="1F2937", sz="16")
            txt_c = RGBColor(31, 41, 55)
        else:
            bg_hex = TTC_BG or "14532D"
            for cell in r.cells:
                _set_cell_bg(cell, bg_hex)
            txt_c = TTC_TXT or WHITE

        for cell, text, align in [(r.cells[0], label, WD_ALIGN_PARAGRAPH.LEFT),
                                   (r.cells[1], value, WD_ALIGN_PARAGRAPH.RIGHT)]:
            p = cell.paragraphs[0]
            p.alignment = align
            _zero_para_spacing(p)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = BODY_FONT
            run.font.color.rgb = txt_c

    def _add_total_net(label: str, value: str):
        """Ligne NET À PAYER."""
        r = tot.add_row()
        if is_pro:
            bg_hex = "F8F9FA"
            for cell in r.cells:
                _set_cell_bg(cell, bg_hex)
                _set_cell_border_bottom(cell, color="1F2937", sz="20")
            txt_c = RGBColor(31, 41, 55)
        else:
            bg_hex = NET_BG or "047857"
            for cell in r.cells:
                _set_cell_bg(cell, bg_hex)
            txt_c = NET_TXT or WHITE

        for cell, text, align in [(r.cells[0], label, WD_ALIGN_PARAGRAPH.LEFT),
                                   (r.cells[1], value, WD_ALIGN_PARAGRAPH.RIGHT)]:
            p = cell.paragraphs[0]
            p.alignment = align
            _zero_para_spacing(p)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = BODY_FONT
            run.font.color.rgb = txt_c

    if with_tva:
        ht_label = "Total HT brut" if has_remise else "Total HT"
        _add_total_row(ht_label, _fmt_money(totaux.total_ht))
        if has_remise:
            ht_net = totaux.total_ht_net if totaux.total_ht_net else (totaux.total_ht - (totaux.remise_ht or 0))
            _add_total_row("Remise", "- " + _fmt_money(totaux.remise_ht or 0))
            _add_total_row("Total HT net", _fmt_money(ht_net))
        _add_total_row("Total TVA", _fmt_money(totaux.total_tva))
        _add_total_accent("TOTAL TTC", _fmt_money(totaux.total_ttc))
        if has_acompte:
            _add_total_row("Acompte verse", "- " + _fmt_money(devis.acompte or 0))
            net = totaux.net_a_payer if totaux.net_a_payer else max(0, totaux.total_ttc - (devis.acompte or 0))
            _add_total_net("NET A PAYER", _fmt_money(net))
    else:
        if has_remise:
            ht_net = totaux.total_ht_net if totaux.total_ht_net else (totaux.total_ht - (totaux.remise_ht or 0))
            _add_total_row("Total HT brut", _fmt_money(totaux.total_ht))
            _add_total_row("Remise", "- " + _fmt_money(totaux.remise_ht or 0))
            _add_total_accent("TOTAL HT NET", _fmt_money(ht_net))
        else:
            _add_total_accent("TOTAL HT", _fmt_money(totaux.total_ht))
        if has_acompte:
            ht_base = totaux.total_ht_net if totaux.total_ht_net else totaux.total_ht
            _add_total_row("Acompte verse", "- " + _fmt_money(devis.acompte or 0))
            net = totaux.net_a_payer if totaux.net_a_payer else max(0, ht_base - (devis.acompte or 0))
            _add_total_net("NET A PAYER", _fmt_money(net))

    sep4 = doc.add_paragraph()
    _zero_para_spacing(sep4)

    # ── MENTIONS LÉGALES ─────────────────────────────────────────────
    sep5 = doc.add_paragraph("─" * 90)
    _zero_para_spacing(sep5)

    for mention in devis.mentions_legales:
        p = doc.add_paragraph(f"• {mention}")
        _zero_para_spacing(p)
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.name = BODY_FONT
        p.runs[0].font.color.rgb = RGBColor(90, 99, 93)

    if not with_tva:
        p = doc.add_paragraph("• TVA non applicable, art. 293 B du CGI")
        _zero_para_spacing(p)
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.name = BODY_FONT
        p.runs[0].font.bold = True

    # ── COORDONNÉES BANCAIRES ─────────────────────────────────────────
    if devis.artisan.iban or devis.artisan.bic:
        sep6 = doc.add_paragraph("─" * 90)
        _zero_para_spacing(sep6)
        p = doc.add_paragraph()
        _zero_para_spacing(p)
        _run(p, "Coordonnees bancaires", size=10, color=COL_ACCENT, bold=True)
        if devis.artisan.iban:
            p2 = doc.add_paragraph(f"IBAN : {devis.artisan.iban}")
            _zero_para_spacing(p2)
            p2.runs[0].font.size = Pt(9)
            p2.runs[0].font.name = BODY_FONT
        if devis.artisan.bic:
            p3 = doc.add_paragraph(f"BIC/SWIFT : {devis.artisan.bic}")
            _zero_para_spacing(p3)
            p3.runs[0].font.size = Pt(9)
            p3.runs[0].font.name = BODY_FONT

    # ── ZONE DE SIGNATURE ────────────────────────────────────────────
    sep_sig = doc.add_paragraph()
    _zero_para_spacing(sep_sig)

    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.autofit = False
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _fix_table(sig_tbl, width_cm=PAGE_W_CM)
    sig_tbl.columns[0].width = Cm(8.0)
    sig_tbl.columns[1].width = Cm(9.0)

    def _sig_cell(cell, title: str, lines: list):
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _el
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = _el("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            b = _el(f"w:{side}")
            b.set(_qn("w:val"), "single")
            b.set(_qn("w:sz"), "4")
            b.set(_qn("w:color"), "CCCCCC")
            tcBorders.append(b)
        tcPr.append(tcBorders)

        p = cell.paragraphs[0]
        _zero_para_spacing(p)
        _run(p, title, size=9, color=COL_ACCENT, bold=True)
        for line in lines:
            pl = cell.add_paragraph()
            _zero_para_spacing(pl)
            _run(pl, line, size=8, color=RGBColor(90, 99, 93))
        for _ in range(3):
            pe = cell.add_paragraph()
            _zero_para_spacing(pe)

    _sig_cell(
        sig_tbl.cell(0, 0),
        "Bon pour accord",
        ["Fait a : _________________________________",
         "Le : _______ / _______ / ___________"],
    )
    _sig_cell(sig_tbl.cell(0, 1), "Signature du client", [])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
