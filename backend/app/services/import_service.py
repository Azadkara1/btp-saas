"""
Import de devis/factures existants (PDF ou .docx).
Envoie le document à Claude pour extraction, puis réutilise l'injection post-génération.
"""
import base64
import io
import json
import logging
import re

import anthropic

from app.core.config import get_settings
from app.core.prompts import IMPORT_EXTRACTION_PROMPT
from app.models.quote import Devis, QuoteRequest, QuoteResponse

settings = get_settings()
_client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

MAX_FILE_BYTES  = 10 * 1024 * 1024  # 10 Mo
MAX_OUTPUT_TOKENS = 8000  # garde-fou : ~50 lignes JSON + descriptions groupées


async def import_quote(
    file_content: bytes,
    filename: str,
    request: QuoteRequest,
) -> QuoteResponse:
    """Point d'entrée : reçoit le fichier brut, dispatche selon l'extension."""
    if len(file_content) > MAX_FILE_BYTES:
        mb = len(file_content) / (1024 * 1024)
        return QuoteResponse(
            success=False,
            error=f"Fichier trop volumineux ({mb:.1f} Mo). Maximum autorisé : 10 Mo.",
        )

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        result = _extract_from_pdf(file_content)
    elif ext in ("docx", "doc"):
        result = _extract_from_docx(file_content)
    else:
        return QuoteResponse(
            success=False,
            error="Format non supporté. Envoyez un fichier PDF ou .docx.",
        )

    if result.success and result.devis:
        _inject_artisan(result.devis, request)

    return result


def _extract_from_pdf(pdf_bytes: bytes) -> QuoteResponse:
    """Envoie le PDF en bloc document natif à Claude (pas de parseur tiers)."""
    pdf_b64 = base64.standard_b64encode(pdf_bytes).decode("utf-8")
    try:
        response = _client.messages.create(
            model=settings.claude_model,
            max_tokens=MAX_OUTPUT_TOKENS,
            system=IMPORT_EXTRACTION_PROMPT,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {
                                "type": "base64",
                                "media_type": "application/pdf",
                                "data": pdf_b64,
                            },
                        },
                        {
                            "type": "text",
                            "text": "Extrais les informations de ce document et retourne le JSON demandé.",
                        },
                    ],
                }
            ],
        )
        tokens = response.usage.input_tokens + response.usage.output_tokens
        logging.info("[IMPORT_PDF] stop=%s tokens=%d", response.stop_reason, tokens)
        return _parse_import_response(response, tokens)
    except Exception as exc:
        logging.error("[IMPORT_PDF] Erreur Claude : %s", exc, exc_info=True)
        return QuoteResponse(success=False, error="Erreur lors de l'analyse du PDF par l'IA.")


def _extract_from_docx(docx_bytes: bytes) -> QuoteResponse:
    """Extrait le texte du .docx via python-docx, puis envoie à Claude."""
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(docx_bytes))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()

        if not text:
            return QuoteResponse(success=False, error="Le fichier .docx semble vide ou illisible.")
    except Exception as exc:
        logging.error("[IMPORT_DOCX] Lecture .docx échouée : %s", exc, exc_info=True)
        return QuoteResponse(success=False, error="Impossible de lire ce fichier .docx.")

    try:
        response = _client.messages.create(
            model=settings.claude_model,
            max_tokens=MAX_OUTPUT_TOKENS,
            system=IMPORT_EXTRACTION_PROMPT,
            messages=[
                {
                    "role": "user",
                    "content": f"Voici le contenu du document à extraire :\n\n{text}",
                }
            ],
        )
        tokens = response.usage.input_tokens + response.usage.output_tokens
        logging.info("[IMPORT_DOCX] stop=%s tokens=%d", response.stop_reason, tokens)
        return _parse_import_response(response, tokens)
    except Exception as exc:
        logging.error("[IMPORT_DOCX] Erreur Claude : %s", exc, exc_info=True)
        return QuoteResponse(success=False, error="Erreur lors de l'analyse du document .docx par l'IA.")


def _parse_import_response(response, total_tokens: int) -> QuoteResponse:
    """
    Parse la réponse Claude dédiée à l'import.
    Détecte les troncatures max_tokens, logue la réponse brute sur échec.
    Ne partage pas l'implémentation avec la génération pour rester indépendant.
    """
    # ── Garde-fou troncature ─────────────────────────────────────
    if response.stop_reason == "max_tokens":
        raw = next((b.text for b in response.content if hasattr(b, "text")), "")
        logging.error(
            "[IMPORT] TRONCATURE — stop_reason=max_tokens (%d tokens). "
            "Fin de réponse reçue : ...%.300s",
            total_tokens, raw[-300:] if raw else "(vide)"
        )
        return QuoteResponse(
            success=False,
            error=(
                "Le document est trop volumineux pour être extrait en une fois. "
                "Essayez avec un document plus court ou comportant moins de lignes."
            ),
        )

    # ── Récupère le bloc texte ───────────────────────────────────
    text_content = next(
        (b.text for b in response.content if hasattr(b, "text")),
        None,
    )
    if not text_content:
        logging.error(
            "[IMPORT] Aucun bloc texte dans la réponse (blocs: %s)",
            [getattr(b, "type", "?") for b in response.content],
        )
        return QuoteResponse(success=False, error="Aucun contenu dans la réponse de l'IA.")

    # ── Strip défensif des fences markdown (```json ... ```) ─────
    clean = text_content.strip()
    if clean.startswith("```"):
        clean = re.sub(r"^```[a-z]*\n?", "", clean)
        clean = re.sub(r"\n?```\s*$", "", clean.strip())
        clean = clean.strip()

    # ── Extraction du bloc JSON ──────────────────────────────────
    match = re.search(r"\{.*\}", clean, re.DOTALL)
    if not match:
        logging.error(
            "[IMPORT] Pas de JSON trouvé. Réponse brute (500 premiers chars) : %.500s",
            text_content,
        )
        return QuoteResponse(success=False, error="L'IA n'a pas retourné de JSON valide.")

    json_text = match.group(0)
    try:
        data = json.loads(json_text)

        # ── Champs enrichis hors schéma Devis → import_meta ─────────
        meta: dict = {}
        for key in ("document_type", "numero_document_original", "date_document_original", "emetteur"):
            meta[key] = data.pop(key, None)

        # Conditions de paiement et acompte : vont dans Devis ET dans meta
        meta["conditions_paiement"] = data.get("conditions_paiement")
        meta["acompte"] = data.get("acompte")

        # numero_document_original → pré-remplit devis.numero_document
        if meta["numero_document_original"] and "numero_document" not in data:
            data["numero_document"] = meta["numero_document_original"]

        # emetteur → utilisé comme artisan dans le Devis (sera écrasé par _inject_artisan)
        if meta["emetteur"] and "artisan" not in data:
            data["artisan"] = {
                "nom": meta["emetteur"].get("nom"),
                "siret": meta["emetteur"].get("siret"),
            }

        devis = Devis(**data)
        return QuoteResponse(success=True, devis=devis, tokens_used=total_tokens, import_meta=meta)
    except json.JSONDecodeError as exc:
        logging.error(
            "[IMPORT] JSON invalide : %s — JSON extrait (500 chars) : %.500s",
            exc, json_text,
        )
        return QuoteResponse(success=False, error="Erreur de format dans la réponse de l'IA.")
    except Exception as exc:
        logging.error("[IMPORT] Erreur validation Pydantic : %s", exc, exc_info=True)
        return QuoteResponse(success=False, error="Erreur lors de la validation du document extrait.")


def _inject_artisan(devis: Devis, request: QuoteRequest) -> None:
    """
    Écrase les infos artisan extraites par celles du profil enregistré.
    Même logique que claude_service.py — l'artisan importé est ignoré.
    """
    a = devis.artisan
    if request.artisan_nom:         a.nom         = request.artisan_nom
    if request.artisan_siret:       a.siret       = request.artisan_siret
    if request.artisan_iban:        a.iban        = request.artisan_iban
    if request.artisan_bic:         a.bic         = request.artisan_bic
    if request.artisan_adresse:     a.adresse     = request.artisan_adresse
    if request.artisan_code_postal: a.code_postal = request.artisan_code_postal
    if request.artisan_ville:       a.ville       = request.artisan_ville
    if request.artisan_telephone:   a.telephone   = request.artisan_telephone
    if request.artisan_email:       a.email       = request.artisan_email
    if request.artisan_site_web:    a.site_web    = request.artisan_site_web
    if request.artisan_logo_base64: a.logo_base64 = request.artisan_logo_base64
    devis.modele = request.modele or "moderne"
